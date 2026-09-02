from pathlib import Path
from types import SimpleNamespace

import pymupdf
import pytest
from docx import Document

from kerui_recruit.resumes import extract as extract_module
from kerui_recruit.resumes.extract import (
    LegacyDocConversionError,
    extract_contact,
    extract_text,
)


def test_pdf_and_docx_extract_visible_text(tmp_path: Path) -> None:
    """A readable office document must not be sent through costly OCR."""
    pdf_path = tmp_path / "resume.pdf"
    pdf = pymupdf.open()
    pdf.new_page().insert_text((72, 72), "Python Finance Resume")
    pdf.save(pdf_path)
    pdf.close()
    docx_path = tmp_path / "resume.docx"
    docx = Document()
    docx.add_paragraph("Python Finance Resume")
    docx.save(docx_path)

    pdf_result = extract_text(pdf_path)
    docx_result = extract_text(docx_path)

    assert "Python Finance" in pdf_result.text
    assert "Python Finance" in docx_result.text
    assert pdf_result.requires_ocr is False
    assert docx_result.requires_ocr is False


def test_blank_pdf_requires_ocr(tmp_path: Path) -> None:
    """A scan with no text layer must never be marked as successfully extracted."""
    pdf_path = tmp_path / "scan.pdf"
    pdf = pymupdf.open()
    pdf.new_page()
    pdf.save(pdf_path)
    pdf.close()

    result = extract_text(pdf_path)

    assert result.text == ""
    assert result.requires_ocr is True
    assert result.page_count == 1


def test_extract_contact_finds_email_and_mainland_mobile() -> None:
    """Sensitive contact details must be discoverable for encrypted persistence."""
    contact = extract_contact(
        "张三\n邮箱: zhang.san@example.com\n手机: 13800138000\n电话: 010-12345678"
    )

    assert contact.email == "zhang.san@example.com"
    assert contact.phone == "13800138000"


def test_extract_contact_returns_none_when_absent() -> None:
    """A resume without contact details must yield no false positives."""
    contact = extract_contact("张三 本科 6年 Java Python")

    assert contact.email is None
    assert contact.phone is None


def test_legacy_doc_extraction_preserves_the_source_file(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "legacy.doc"
    original = b"legacy-binary-content"
    source.write_bytes(original)
    monkeypatch.setattr(
        extract_module,
        "_extract_legacy_doc",
        lambda _path: "Python Finance Resume",
    )

    result = extract_text(source)

    assert result.text == "Python Finance Resume"
    assert result.requires_ocr is False
    assert source.read_bytes() == original


def test_windows_legacy_doc_requires_microsoft_word(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"legacy")
    monkeypatch.setattr(extract_module.platform, "system", lambda: "Windows")

    def missing_module(_name: str):
        raise ModuleNotFoundError

    monkeypatch.setattr(extract_module.importlib, "import_module", missing_module)

    with pytest.raises(LegacyDocConversionError, match="Microsoft Word.*DOCX"):
        extract_module._extract_legacy_doc(source)


def test_macos_legacy_doc_uses_system_textutil(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"legacy")
    captured: list[str] = []

    def fake_run(command: list[str], **_kwargs):
        captured.extend(command)
        return SimpleNamespace(returncode=0, stdout="Python Resume", stderr="")

    monkeypatch.setattr(extract_module.platform, "system", lambda: "Darwin")
    monkeypatch.setattr(extract_module.subprocess, "run", fake_run)

    assert extract_module._extract_legacy_doc(source) == "Python Resume"
    assert captured == ["textutil", "-convert", "txt", "-stdout", str(source.resolve())]


def test_windows_word_converter_initializes_com_and_never_saves(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "legacy.doc"
    source.write_bytes(b"legacy")
    calls: list[object] = []

    class FakeDocument:
        Content = SimpleNamespace(Text="Python Resume")

        def Close(self, *, SaveChanges: int) -> None:
            calls.append(("close", SaveChanges))

    class FakeWord:
        Documents = SimpleNamespace(Open=lambda *_args, **_kwargs: FakeDocument())
        Visible = True
        DisplayAlerts = 1

        def Quit(self) -> None:
            calls.append("quit")

    fake_pythoncom = SimpleNamespace(
        CoInitialize=lambda: calls.append("initialize"),
        CoUninitialize=lambda: calls.append("uninitialize"),
    )
    fake_client = SimpleNamespace(DispatchEx=lambda _name: FakeWord())

    def fake_import(name: str):
        return fake_pythoncom if name == "pythoncom" else fake_client

    monkeypatch.setattr(extract_module.importlib, "import_module", fake_import)

    assert extract_module._extract_doc_with_word(source) == "Python Resume"
    assert calls == ["initialize", ("close", 0), "quit", "uninitialize"]


def _image_pixmap(width: int, height: int) -> pymupdf.Pixmap:
    pixmap = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, width, height))
    pixmap.clear_with(0)
    return pixmap


def test_image_body_with_repeated_watermark_requires_ocr(tmp_path: Path) -> None:
    """图片正文 + 超过 20 字符的重复水印，必须进入 OCR。"""
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_image(page.rect, pixmap=_image_pixmap(600, 800))
    watermark = "机密-内部资料-禁止外传-CONFIDENTIAL"
    for offset in range(8):
        page.insert_text((72, 72 + offset * 20), watermark)
    path = tmp_path / "watermark.pdf"
    pdf.save(path)
    pdf.close()

    result = extract_text(path)

    assert result.requires_ocr is True
    assert result.pages[0].needs_ocr is True
    assert "水印" in result.pages[0].reason


def test_normal_chinese_text_pdf_is_not_ocred(tmp_path: Path) -> None:
    pdf = pymupdf.open()
    page = pdf.new_page()
    lines = [
        "张三",
        "求职意向：后端工程师（风控）",
        "工作经历：某科技公司，负责支付风控平台研发，使用 Java 与 Python 构建高并发交易系统。",
        "教育经历：某大学 计算机科学与技术 本科 2016-2020",
        "技能：Java、Python、MySQL、Redis、微服务",
    ]
    for offset, line in enumerate(lines):
        page.insert_text((72, 72 + offset * 18), line, fontname="china-s")
    path = tmp_path / "zh.pdf"
    pdf.save(path)
    pdf.close()

    result = extract_text(path)

    assert result.requires_ocr is False
    assert result.pages[0].needs_ocr is False


def test_english_resume_is_not_ocred_for_low_cjk(tmp_path: Path) -> None:
    pdf = pymupdf.open()
    page = pdf.new_page()
    lines = [
        "John Smith",
        "Senior Software Engineer - Payments & Risk",
        "Built real-time fraud detection services using Python and Java.",
        "Education: Bachelor of Computer Science, 2016-2020",
        "Skills: Python, Java, MySQL, Redis, Microservices",
    ]
    for offset, line in enumerate(lines):
        page.insert_text((72, 72 + offset * 18), line)
    path = tmp_path / "en.pdf"
    pdf.save(path)
    pdf.close()

    result = extract_text(path)

    assert result.requires_ocr is False
    assert result.pages[0].needs_ocr is False


def test_text_resume_with_logo_is_not_ocred(tmp_path: Path) -> None:
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text(
        (72, 72),
        "张三\n工作经历：某公司 风控工程师 5年\n技能：Python 风控 支付",
        fontname="china-s",
    )
    # 角落放一个小 Logo，不应导致整页 OCR。
    page.insert_image(
        pymupdf.Rect(500, 700, 560, 760),
        pixmap=_image_pixmap(60, 60),
    )
    path = tmp_path / "logo.pdf"
    pdf.save(path)
    pdf.close()

    result = extract_text(path)

    assert result.requires_ocr is False
    assert result.pages[0].needs_ocr is False


def test_mixed_pdf_only_ocred_the_image_page(tmp_path: Path) -> None:
    pdf = pymupdf.open()
    text_page = pdf.new_page()
    text_page.insert_text((72, 72), "张三\n风控工程师\nPython 支付风控", fontname="china-s")
    image_page = pdf.new_page()
    image_page.insert_image(image_page.rect, pixmap=_image_pixmap(600, 800))
    path = tmp_path / "mixed.pdf"
    pdf.save(path)
    pdf.close()

    result = extract_text(path)

    assert result.pages[0].needs_ocr is False
    assert result.pages[1].needs_ocr is True
    assert result.requires_ocr is True

