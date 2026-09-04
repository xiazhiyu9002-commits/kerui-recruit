from __future__ import annotations

import re
from io import BytesIO
from typing import Protocol

from docx import Document

from kerui_recruit.org.structured import OrgClarificationQuestion, OrgParseResult, ParsedOrgDraft
from kerui_recruit.providers.openai_compatible import OpenAICompatibleClient


class OrgImportParser(Protocol):
    async def parse(self, text: str) -> ParsedOrgDraft: ...

    async def parse_clarifying(self, text: str, answers: list[str] | None = None) -> OrgParseResult: ...

    async def parse_incremental(self, text: str, answers: list[str] | None = None) -> OrgParseResult: ...

    async def revise(self, draft: ParsedOrgDraft, instruction: str) -> ParsedOrgDraft: ...


def extract_org_text(filename: str, content: bytes) -> str:
    """Extract plain text from a .txt or .docx org description."""
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if suffix == "docx":
        document = Document(BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend(cell.text for row in table.rows for cell in row.cells)
        return "\n".join(line.strip() for line in "\n".join(parts).splitlines() if line.strip())
    if suffix == "txt":
        return content.decode("utf-8", errors="replace")
    raise ValueError("仅支持 .txt 或 .docx 文件")


_SENTENCE_ENDINGS = ("。", "；", "！", "？")
_INC_USE_THRESHOLD = 3000  # 全文超过该字数走增量解析


def _chunk_text(text: str, *, max_chars: int = 1500) -> list[str]:
    """按语义边界切分长文本：优先按空行分段；单段超长再按句末标点切分，绝不切断半句。"""
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    chunks: list[str] = []
    for paragraph in paragraphs:
        if len(paragraph) <= max_chars:
            chunks.append(paragraph)
            continue
        buffer = ""
        for char in paragraph:
            buffer += char
            if char in _SENTENCE_ENDINGS and len(buffer) >= max_chars:
                chunks.append(buffer)
                buffer = ""
        if buffer.strip():
            chunks.append(buffer)
    return chunks


_ORG_PARSE_PROMPT = """你是组织架构整理助手，负责把一段关于公司组织/团队/人员的非结构化中文描述，解析成结构化 JSON。

任务：根据原文输出一个 JSON 对象，字段如下：
- company_name：公司名称（如「得物」或原文提到的公司；若原文只描述团队，用最能概括的名称，无法判断填「未命名公司」）；
- departments：部门数组，每项含：
  - name：部门/团队名称；
  - parent_name：上级部门名称（无上级填 null）；
  - leader_name：该部门负责人姓名（可空）；
  - team_size：团队人数数字（如「70 来人」取 70；「百人」取 100；无法判断填 null）；
  - business_direction：该部门业务方向一句话（可空）；
- employees：人员数组，每项含：
  - name：本名（若只出现花名，用花名作为 name 并在 alias 留空）；
  - alias：花名/别名（如「叶程」的花名对应本名「贺喜」时，name=贺喜、alias=叶程；反之亦然；没有别名填 null）；
  - title：职位/头衔（如「算法平台负责人」「社区推荐 TL」）；
  - job_level：职级（如「字节 4-1」「P8」，无法判断填 null）；
  - report_to_name：直接汇报给谁（姓名），无汇报对象填 null；
  - department_name：所属部门/团队名称（可空）；
  - subordinate_count：下属人数数字（「下属 70 来人」取 70；「三百来号人」取 300；无法判断填 null）；
  - team_size：该人带的小团队人数（与部门团队规模不同，无法判断填 null）；
  - remark：背景/备注，把该人的来源、履历、风格、负责范围等原文信息概括成一句话（可空）。

规则：
- 姓名优先用本名，花名放到 alias；若原文只给了花名，则 name 用花名；
- 「汇报给 X」→ report_to_name=X；「X 平行于 Y」→ 都汇报给同一上级；
- 人数、职级、背景尽量从原文提取，无法判断填 null；
- 只输出 JSON 对象，不要 markdown 代码块或任何多余文字。

原文：
{org_text}"""


_ORG_REVISE_PROMPT = """你是组织架构整理助手。下面是一份已解析的组织架构草稿（JSON）与用户的一条修正指令。请根据指令修改草稿，并返回修改后的完整 JSON。

草稿：
{draft}

修正指令：
{instruction}

规则：
- 只根据指令做针对性修改，其余字段保持原样；
- 保持与草稿相同的 JSON 结构（company_name / departments / employees）；
- 只输出 JSON 对象，不要 markdown 代码块或任何多余文字。"""


_ORG_CLARIFY_PROMPT = """你是组织架构整理助手，负责把一段关于公司组织/团队/人员的非结构化中文描述，解析成结构化 JSON。如果你对某些关键信息不明确或存在歧义，不要擅自编造，而是把疑问列出来向用户提问。

输出一个 JSON 对象，结构如下：
- draft：一个对象或 null，字段同下：
  - company_name：公司名称（无法判断填「未命名公司」）；
  - departments：部门数组，每项含 name / parent_name / leader_name / team_size / business_direction；
  - employees：人员数组，每项含 name / alias / title / job_level / report_to_name / department_name / subordinate_count / team_size / remark；
- questions：疑问数组（可为空），每项含：
  - question：向用户提出的问题（中文，明确、具体）；
  - field：该问题关联的字段名（如 company_name / department_name / report_to_name / team_size，无法对应填 null）；
  - hint：可选的补充说明（可空）。

规则：
- 姓名优先用本名，花名放到 alias；只给花名则 name 用花名；
- 人数、职级、背景尽量从原文提取，无法判断填 null；
- 信息足够时，draft 正常填充，questions 为空数组；
- 信息不足时，draft 中能确定的字段尽量填充，其余疑问放入 questions；
- 如果用户已提供补充回答，优先依据回答修正 draft，不要再重复提问；
- 只输出 JSON 对象，不要 markdown 代码块或任何多余文字。

{answers_block}

原文：
{org_text}"""


_ORG_INCREMENTAL_PROMPT = """你是组织架构整理助手。我会分段给你组织描述文本，你需要把每段信息增量合并进「当前草稿」，并输出更新后的完整草稿。

输出一个 JSON 对象，结构如下：
- draft：更新后的完整草稿，字段同下：
  - company_name：公司名称（无法判断填「未命名公司」）；
  - departments：部门数组，每项含 name / parent_name / leader_name / team_size / business_direction；
  - employees：人员数组，每项含 name / alias / title / job_level / report_to_name / department_name / subordinate_count / team_size / remark；
- questions：疑问数组（可为空），每项含：
  - question：向用户提出的问题（中文，明确、具体）；
  - field：该问题关联的字段名（无法对应填 null）；
  - hint：可选的补充说明（可空）。

合并规则：
- 把本段新信息合并进当前草稿，不得丢弃已有信息；
- 同名或花名/本名指向同一人时合并（如「叶程」与「贺喜」是同一个人）；
- 跨段补全汇报关系（上级在前段、下属在后段时回填 report_to_name）；
- 前后段对同一人描述冲突时，取更具体或更新的信息；无法判定则写入 questions；
- 人数、职级、背景尽量从原文提取，无法判断填 null；
- 只输出 JSON 对象，不要 markdown 代码块或任何多余文字。

当前草稿：
{current_draft}

本段原文：
{chunk_text}"""


class DeepSeekOrgImportParser:
    """Parse unstructured org text into :class:`ParsedOrgDraft`."""

    def __init__(
        self,
        *,
        api_key: str,
        client,
        base_url: str = "https://api.deepseek.com",
        model: str = "deepseek-v4-flash",
    ) -> None:
        self._llm = OpenAICompatibleClient(
            base_url=base_url,
            api_key=api_key,
            model=model,
            http_client=client,
        )

    async def parse(self, text: str) -> ParsedOrgDraft:
        return await self._llm.complete_json(
            messages=[
                {"role": "user", "content": _ORG_PARSE_PROMPT.format(org_text=text)}
            ],
            response_model=ParsedOrgDraft,
        )

    async def parse_clarifying(self, text: str, answers: list[str] | None = None) -> OrgParseResult:
        if len(text) > _INC_USE_THRESHOLD:
            return await self.parse_incremental(text, answers)
        answers_block = ""
        if answers:
            answers_block = "用户对疑问的补充回答：\n" + "\n".join(f"- {answer}" for answer in answers)
        return await self._llm.complete_json(
            messages=[
                {
                    "role": "user",
                    "content": _ORG_CLARIFY_PROMPT.format(org_text=text, answers_block=answers_block),
                }
            ],
            response_model=OrgParseResult,
        )

    async def parse_incremental(self, text: str, answers: list[str] | None = None) -> OrgParseResult:
        chunks = _chunk_text(text)
        draft = ParsedOrgDraft(company_name="未命名公司")
        questions: list[OrgClarificationQuestion] = []
        for chunk in chunks:
            result = await self._llm.complete_json(
                messages=[
                    {
                        "role": "user",
                        "content": _ORG_INCREMENTAL_PROMPT.format(
                            current_draft=draft.model_dump_json(),
                            chunk_text=chunk,
                        ),
                    }
                ],
                response_model=OrgParseResult,
            )
            if result.draft is not None:
                draft = result.draft
            questions.extend(result.questions)
        return OrgParseResult(draft=draft, questions=questions)

    async def revise(self, draft: ParsedOrgDraft, instruction: str) -> ParsedOrgDraft:
        return await self._llm.complete_json(
            messages=[
                {
                    "role": "user",
                    "content": _ORG_REVISE_PROMPT.format(
                        draft=draft.model_dump_json(),
                        instruction=instruction,
                    ),
                }
            ],
            response_model=ParsedOrgDraft,
        )
