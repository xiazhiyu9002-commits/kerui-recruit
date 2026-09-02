from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from openpyxl import load_workbook


class UnsupportedJdType(ValueError):
    code = "E_JD_FILE_TYPE_UNSUPPORTED"


_JD_HEADER_RE = re.compile(
    r"^\s*(?:岗位|职位|招聘岗位|招聘职位|JD|Job)\s*[0-9一二三四五六七八九十]+"
)


def split_jd_text(text: str) -> list[str]:
    """Split raw JD text into individual JD chunks by header-like lines.

    Because users are free to paste multiple JDs in one blob, split on lines
    such as ``岗位1`` / ``职位二`` / ``JD 2`` and fall back to the whole text
    when no such headers are found.
    """
    lines = text.splitlines()
    chunks: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if _JD_HEADER_RE.match(line) and current:
            chunks.append(current)
            current = [line]
        else:
            current.append(line)
    if current:
        chunks.append(current)

    result = ["\n".join(chunk).strip() for chunk in chunks]
    result = [chunk for chunk in result if chunk]
    return result or [text.strip()]


def extract_jd_text(filename: str, content: bytes) -> str:
    """Extract plain text from a Word or Excel JD file."""
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if suffix == "docx":
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend(cell.text for row in table.rows for cell in row.cells)
        return _normalize("\n".join(paragraphs))
    if suffix == "xlsx":
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        rows: list[str] = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                rows.append(" ".join(str(cell) for cell in row if cell is not None))
        return _normalize("\n".join(rows))
    raise UnsupportedJdType(f"Unsupported JD type: {suffix or 'none'}")


def _normalize(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())
