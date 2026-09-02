from __future__ import annotations

import importlib
import platform
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from docx import Document

from kerui_recruit.resumes.quality import (
    DOMINANT_FRAGMENT_RATIO,
    IMAGE_HEAVY_COVERAGE,
    MIN_MEANINGFUL_CHARS,
    analyze_text,
)

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"(?<!\d)1[3-9](?:[\s-]?\d){9}(?!\d)")


class LegacyDocConversionError(RuntimeError):
    code = "E_DOC_CONVERTER_UNAVAILABLE"


@dataclass(frozen=True, slots=True)
class ExtractedContact:
    email: str | None
    phone: str | None


def extract_contact(text: str) -> ExtractedContact:
    """Extract a best-effort email and mainland mobile number from raw text."""
    email_match = _EMAIL_RE.search(text)
    phone_match = _PHONE_RE.search(text)
    phone = phone_match.group(0) if phone_match else None
    if phone:
        phone = re.sub(r"\D", "", phone)
    return ExtractedContact(
        email=email_match.group(0) if email_match else None,
        phone=phone,
    )


@dataclass(frozen=True, slots=True)
class PageAssessment:
    page_index: int
    text: str
    needs_ocr: bool
    reason: str
    valid_char_count: int
    repeated_ratio: float
    dominant_ratio: float
    image_coverage: float


@dataclass(frozen=True, slots=True)
class ExtractedText:
    text: str
    page_count: int
    requires_ocr: bool
    pages: tuple[PageAssessment, ...] = ()


def extract_text(path: Path) -> ExtractedText:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        with pymupdf.open(path) as document:
            pages = tuple(
                _assess_pdf_page(document[index], index)
                for index in range(document.page_count)
            )
        text = _normalize_text("\n".join(page.text for page in pages))
        return ExtractedText(
            text=text,
            page_count=len(pages),
            requires_ocr=any(page.needs_ocr for page in pages),
            pages=pages,
        )
    if suffix == ".docx":
        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend(cell.text for row in table.rows for cell in row.cells)
        return ExtractedText(
            text=_normalize_text("\n".join(paragraphs)),
            page_count=1,
            requires_ocr=False,
        )
    if suffix == ".doc":
        text = _normalize_text(_extract_legacy_doc(path))
        if not text:
            raise LegacyDocConversionError(
                "旧版 DOC 未提取到文字，请将文件另存为 DOCX 后重试"
            )
        return ExtractedText(text=text, page_count=1, requires_ocr=False)
    raise ValueError(f"Text extraction is unsupported for {suffix or 'unknown'}")


def _extract_legacy_doc(path: Path) -> str:
    system = platform.system()
    if system == "Windows":
        return _extract_doc_with_word(path)
    if system == "Darwin":
        return _extract_doc_with_textutil(path)
    raise LegacyDocConversionError(
        "当前系统不支持旧版 DOC，请将文件另存为 DOCX 后重试"
    )


def _extract_doc_with_word(path: Path) -> str:
    try:
        pythoncom = importlib.import_module("pythoncom")
        win32_client = importlib.import_module("win32com.client")
    except ModuleNotFoundError as error:
        raise LegacyDocConversionError(
            "Windows 解析旧版 DOC 需要 Microsoft Word，请安装 Word 或另存为 DOCX 后重试"
        ) from error

    word = None
    document = None
    com_initialized = False
    try:
        pythoncom.CoInitialize()
        com_initialized = True
        word = win32_client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(path.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        return str(document.Content.Text)
    except Exception as error:
        raise LegacyDocConversionError(
            "Microsoft Word 无法读取旧版 DOC，请将文件另存为 DOCX 后重试"
        ) from error
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=0)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if com_initialized:
            pythoncom.CoUninitialize()


def convert_doc_to_pdf(source: Path) -> Path:
    """Convert a .doc/.docx to PDF via Microsoft Word for in-browser preview."""
    import tempfile
    import uuid

    try:
        pythoncom = importlib.import_module("pythoncom")
        win32_client = importlib.import_module("win32com.client")
    except ModuleNotFoundError as error:
        raise LegacyDocConversionError(
            "预览 Word 文档需要安装 Microsoft Word，请安装后重试"
        ) from error

    target = Path(tempfile.gettempdir()) / f"kerui_preview_{uuid.uuid4().hex}.pdf"
    word = None
    document = None
    com_initialized = False
    try:
        pythoncom.CoInitialize()
        com_initialized = True
        word = win32_client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(source.resolve()),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        document.SaveAs2(str(target), FileFormat=17)  # wdFormatPDF
        document.Close(SaveChanges=0)
        document = None
        return target
    except Exception as error:
        raise LegacyDocConversionError(
            "Microsoft Word 无法将文档转换为 PDF 用于预览"
        ) from error
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=0)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if com_initialized:
            pythoncom.CoUninitialize()


def _extract_doc_with_textutil(path: Path) -> str:
    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path.resolve())],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60,
            check=False,
        )
    except (FileNotFoundError, subprocess.TimeoutExpired) as error:
        raise LegacyDocConversionError(
            "macOS textutil 无法使用，请将旧版 DOC 另存为 DOCX 后重试"
        ) from error
    if result.returncode != 0:
        raise LegacyDocConversionError(
            "macOS textutil 无法读取旧版 DOC，请将文件另存为 DOCX 后重试"
        )
    return result.stdout


def _normalize_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())


def _assess_pdf_page(page: pymupdf.Page, page_index: int) -> PageAssessment:
    """逐页判断直接提取的文本是否足以代表页面正文。

    仅把「空白/扫描、水印主导、重复乱码、图片为主且文字过少」的页面判为需要
    OCR，普通文本页（即使带头像、Logo、二维码或中文字符偏少）保持直接提取。
    """
    raw_text = page.get_text("text")
    text = _normalize_text(raw_text)
    quality = analyze_text(text)
    image_coverage = _image_coverage(page)

    needs_ocr = False
    reasons: list[str] = []
    if quality.valid_char_count == 0:
        needs_ocr = True
        reasons.append("页面无可提取文字")
    elif quality.dominant_ratio >= DOMINANT_FRAGMENT_RATIO:
        needs_ocr = True
        reasons.append(f"文字以重复水印为主（占 {quality.dominant_ratio:.0%}）")
    elif (
        quality.valid_char_count < MIN_MEANINGFUL_CHARS
        and image_coverage >= IMAGE_HEAVY_COVERAGE
    ):
        needs_ocr = True
        reasons.append(
            f"页面以图片为主且可提取文字过少（{quality.valid_char_count} 字符）"
        )

    return PageAssessment(
        page_index=page_index,
        text=text,
        needs_ocr=needs_ocr,
        reason="；".join(reasons) if needs_ocr else "文本正常，直接提取",
        valid_char_count=quality.valid_char_count,
        repeated_ratio=quality.repeated_ratio,
        dominant_ratio=quality.dominant_ratio,
        image_coverage=image_coverage,
    )


def _image_coverage(page: pymupdf.Page) -> float:
    """估算页面被图片覆盖的比例（按面积累加、上限 1.0）。"""
    rect = page.rect
    area = rect.width * rect.height
    if area <= 0:
        return 0.0
    infos = page.get_image_info()
    if not infos:
        return 0.0
    covered = 0.0
    for info in infos:
        x0, y0, x1, y1 = info["bbox"]
        width = max(0.0, min(x1, rect.x1) - max(x0, rect.x0))
        height = max(0.0, min(y1, rect.y1) - max(y0, rect.y0))
        covered += width * height
    return min(1.0, covered / area)
